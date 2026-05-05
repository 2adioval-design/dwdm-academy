#!/usr/bin/env python3
"""
extract_catalog.py
Reads DWDM Architect Program .docx files and outputs:
  content/catalog.json  — used by the seed script
  content/catalog.js    — legacy browser-readable export
Usage:
  python scripts/extract_catalog.py [--docs-dir PATH]
  Default docs-dir: ~/Downloads
"""
import re, json, sys, os, argparse
from pathlib import Path

ROOT = Path(__file__).parent.parent
CONTENT_DIR = ROOT / "content"
CONTENT_DIR.mkdir(exist_ok=True)

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx not installed — installing...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.oxml.ns import qn

# ── helpers ──────────────────────────────────────────────────────────────────
def slugify(text):
    s = re.sub(r'[^\w\s-]', '', text.lower())
    return re.sub(r'[\s_]+', '-', s).strip('-')[:80]

def get_outline_level(para):
    """Return outline level (1=Heading1, 2=Heading2, 0=body text)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return 0
    # Try numPr / outlineLvl
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is not None:
        val = outlineLvl.get(qn('w:val'))
        if val is not None:
            return int(val) + 1
    # Try style id
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is not None:
        sid = pStyle.get(qn('w:val'), '')
        m = re.match(r'[Hh]eading(\d+)', sid)
        if m:
            return int(m.group(1))
        # Word sometimes names them Heading1, heading1, or Ttulo1 (Spanish)
    return 0

def table_to_md(table):
    rows = table.rows
    if not rows:
        return ''
    lines = []
    header = [c.text.strip().replace('|', ' ') for c in rows[0].cells]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('|' + '---|' * len(header))
    for row in rows[1:]:
        cells = [c.text.strip().replace('\n', ' ').replace('|', ' ') for c in row.cells]
        lines.append('| ' + ' | '.join(cells) + ' |')
    return '\n'.join(lines)

def para_to_md(para):
    text = para.text.strip()
    if not text:
        return ''
    result = ''
    for run in para.runs:
        t = run.text
        if run.bold and run.italic:
            t = f'***{t}***'
        elif run.bold:
            t = f'**{t}**'
        elif run.italic:
            t = f'*{t}*'
        result += t
    return result.strip() or text

def docx_to_sections(doc_path):
    doc = Document(str(doc_path))
    sections = []
    current = None
    body = doc.element.body

    for block in body:
        tag = block.tag.split('}')[-1]
        if tag == 'p':
            para = None
            for p in doc.paragraphs:
                if p._element is block:
                    para = p
                    break
            if para is None:
                continue
            lvl = get_outline_level(para)
            text = para.text.strip()
            if not text:
                if current:
                    current['lines'].append('')
                continue
            if lvl == 1:
                if current:
                    sections.append(current)
                current = {'title': text, 'slug': slugify(text), 'lines': [f'# {text}', '']}
            elif lvl > 1 and current:
                current['lines'].append('#' * min(lvl, 4) + ' ' + text)
            elif current:
                # body text
                style_name = para.style.name if para.style else ''
                if 'List' in style_name:
                    current['lines'].append(f'- {text}')
                else:
                    current['lines'].append(para_to_md(para))
        elif tag == 'tbl':
            for t in doc.tables:
                if t._element is block:
                    if current:
                        current['lines'].extend(['', table_to_md(t), ''])
                    break

    if current:
        sections.append(current)

    result = []
    for sec in sections:
        md = '\n'.join(sec['lines']).strip()
        md = re.sub(r'\n{3,}', '\n\n', md)
        result.append({'title': sec['title'], 'slug': sec['slug'], 'content_md': md})
    return result

# ── Embedded fallback catalog ─────────────────────────────────────────────────
FALLBACK = {
    1: [
        {"title": "Ch 1 — Optical Networking Landscape", "slug": "ch1-optical-networking-landscape", "content_md": "# Ch 1 — Optical Networking Landscape\n\n## Evolution of Optical Transport\n\n| Era | Technology | Capacity |\n|---|---|---|\n| 1980s | SONET/SDH | 2.5–10 Gbps |\n| 1995–2005 | DWDM fixed-grid | 1.6–4 Tbps |\n| 2010+ | Coherent + flex-grid ROADM | 20–96 Tbps |\n\n## DWDM vs CWDM vs ROADM\n\n| Feature | CWDM | DWDM | ROADM |\n|---|---|---|---|\n| Channel spacing | 20 nm | 50/100 GHz | Flexible |\n| Channel count | 18 max | 96+ | 96+ |\n| Amplification | Passive only | EDFA | EDFA + WSS |"},
        {"title": "Ch 2 — Fiber Fundamentals", "slug": "ch2-fiber-fundamentals", "content_md": "# Ch 2 — Fiber Fundamentals\n\n## Span Loss\n\n```\nSpan Loss = L × α + (n_conn × IL_conn) + (n_splice × IL_splice)\n```\n\n## Fiber Types\n\n| ITU-T | Name | α (dB/km) | CD (ps/nm·km) |\n|---|---|---|---|\n| G.652D | Standard SMF | 0.20 | +17 |\n| G.654E | Ultra-low loss | 0.16 | +21 |\n| G.655C | NZDSF | 0.20 | +4 to +8 |"},
        {"title": "Ch 3 — DWDM Components", "slug": "ch3-dwdm-components", "content_md": "# Ch 3 — DWDM Components\n\n## EDFA Parameters\n\n| Parameter | Typical | Impact |\n|---|---|---|\n| Gain | 15–30 dB | Sets reach |\n| Noise Figure | 4–6 dB | Limits OSNR |\n| Output power | +17 to +23 dBm | Channel count limit |\n\n## WSS (Wavelength Selective Switch)\n\nUses LCoS arrays to route individual wavelengths. Key specs: port count 1×9 to 1×20, insertion loss 4–7 dB, isolation >35 dB."},
        {"title": "Ch 4 — Channel Planning ITU-T G.694.1", "slug": "ch4-channel-planning", "content_md": "# Ch 4 — Channel Planning ITU-T G.694.1\n\n## Reference Grid\n\nReference frequency: **193.1 THz** (1552.52 nm). Spacings: 12.5, 25, 50, 100 GHz.\n\n## C-Band Boundaries\n\n| Boundary | Frequency | Wavelength |\n|---|---|---|\n| Short edge | 196.1 THz | 1528.77 nm |\n| Long edge | 191.7 THz | 1563.86 nm |"},
        {"title": "Ch 5 — Link Budget Calculation", "slug": "ch5-link-budget", "content_md": "# Ch 5 — Link Budget Calculation\n\n## Span Loss Formula\n\n```\nSpan Loss = (L × α) + (n_conn × IL_conn) + margin\n```\n\n## OSNR at Receiver\n\n```\nOSNR = P_out − NF − 10·log10(h·f·Bref) − 10·log10(N_spans)\n```\n\nMinimum required: 14 dB (100G DP-QPSK), 18 dB (200G DP-16QAM)."},
        {"title": "Ch 6 — Modulation Formats & Coherent Optics", "slug": "ch6-modulation-coherent", "content_md": "# Ch 6 — Modulation Formats & Coherent Optics\n\n| Format | Bits/Symbol | Min OSNR | Reach |\n|---|---|---|---|\n| DP-QPSK | 4 | 14 dB | 3000+ km |\n| DP-16QAM | 8 | 22 dB | 800 km |\n| DP-64QAM | 12 | 28 dB | 200 km |\n\n## FEC Types\n\n- **CFEC**: 7% overhead, 8.5 dB NCG\n- **EFEC**: 20% overhead, 11 dB NCG\n- **OFEC**: OpenROADM standardized"},
        {"title": "Ch 7 — Network Topologies & Protection", "slug": "ch7-topologies-protection", "content_md": "# Ch 7 — Network Topologies & Protection\n\n| Topology | Resiliency | Cost |\n|---|---|---|\n| Point-to-Point | None | Low |\n| Ring (OADM) | 1+1 auto | Medium |\n| Mesh (ROADM) | Restoration | High |\n\n## Protection Schemes\n\n- **1+1 optical**: simultaneous Tx both paths, switch <50 ms\n- **1:1 optical**: protection dark, switch <50 ms\n- **GMPLS restoration**: route recalc, 200–500 ms"},
        {"title": "Ch 8 — Operations & Network Management", "slug": "ch8-operations-management", "content_md": "# Ch 8 — Operations & Network Management\n\n## Management Layers\n\n| Layer | Function | Examples |\n|---|---|---|\n| Element (EML) | NE config | OneControl, NetAct |\n| Network (NML) | E2E service | Cisco NSO, MCP |\n| Service (SML) | Customer SLA | OSS/BSS |\n\n## Key PM Thresholds\n\n- OPR: ±1 dB from baseline\n- OSNR: > required + 3 dB\n- Q-factor: > 8.5 dB"},
        {"title": "Ch 9 — Standards & Interoperability", "slug": "ch9-standards-interop", "content_md": "# Ch 9 — Standards & Interoperability\n\n| Body | Focus | Key Standards |\n|---|---|---|\n| ITU-T | Physical | G.694.1, G.652, G.709 |\n| OIF | Interfaces | OIF-400ZR, OIF-CFP2-ACO |\n| OpenROADM | Multi-vendor | Device model, FEC |\n| IETF | Control plane | PCEP, RESTCONF, YANG |"},
        {"title": "Lab 1.1 — Architecture Technology Selection", "slug": "lab-1-1-tech-selection", "content_md": "# Lab 1.1 — Architecture Technology Selection\n\n## Scenario\n\n120 km regional network, 8 wavelengths at 10G today, growing to 40 wavelengths in 3 years.\n\n## Decision\n\nDWDM required: 120 km exceeds passive CWDM reach (~80 km), and 40-channel growth exceeds CWDM capacity."},
        {"title": "Lab 1.6 — Full Link Budget CO-SEA to CO-PDX", "slug": "lab-1-6-link-budget-sea-pdx", "content_md": "# Lab 1.6 — Full Link Budget CO-SEA→CO-PDX\n\n| Parameter | Value |\n|---|---|\n| Route | Seattle → Portland |\n| Distance | 280 km |\n| Spans | 3 (90/95/95 km) |\n| Fiber | G.652D |\n\n## Result\n\nReceived OSNR ≈ 21 dB → **PASS** for 100G DP-QPSK (min 14 dB)"}
    ],
    2: [
        {"title": "Ch 1 — ROADM Architecture Deep Dive", "slug": "ch1-roadm-architecture", "content_md": "# Ch 1 — ROADM Architecture Deep Dive\n\n## CDC-F Architecture\n\nColorless, Directionless, Contentionless, Flexible-grid (CDC-F) enables any wavelength to any direction without recabling.\n\n| Degree Count | WSS Port Count | Max Channels |\n|---|---|---|\n| 2 | 1×9 | 96 |\n| 4 | 1×9 | 96 |\n| 8 | 1×20 | 96 |"},
        {"title": "Ch 2 — Advanced OSNR Engineering", "slug": "ch2-advanced-osnr", "content_md": "# Ch 2 — Advanced OSNR Engineering\n\n## Full Budget Formula\n\n```\nOSNR = P_launch − NF − 10·log10(h·ν·B_ref) − 10·log10(N)\n```\n\n## 1200 km Example (16×75 km spans)\n\n| Parameter | Value |\n|---|---|\n| Linear OSNR | 22.5 dB |\n| NLI penalty | −1.5 dB |\n| Net OSNR | **21.0 dB** |\n| Required (200G) | 18 dB |\n| Margin | **3.0 dB ✓** |"},
        {"title": "Ch 3 — Coherent Transmission Planning", "slug": "ch3-coherent-planning", "content_md": "# Ch 3 — Coherent Transmission Planning\n\n| Format | Line Rate | Reach |\n|---|---|---|\n| DP-QPSK | 100G | 3000+ km |\n| DP-8QAM | 150G | 1800 km |\n| DP-16QAM | 400G | 600 km |\n| DP-64QAM | 800G | 120 km |\n\n## Probabilistic Constellation Shaping (PCS)\n\nApplies Maxwell-Boltzmann distribution, achieving 0.5–1.5 dB SNR gain.\n\n## Rate-Adaptive Design\n\nModern transponders (Ciena WaveLogic 5, Nokia PSE-3) auto-select format based on real-time OSNR."}
    ],
    3: [
        {"title": "Ch 1 — Production Operations Mindset", "slug": "ch1-production-mindset", "content_md": "# Ch 1 — Production Operations Mindset\n\n## Three Commandments\n\n1. **Document everything** — if it isn't in the ticket, it didn't happen\n2. **Follow the MOP** — every change requires a Method of Procedure\n3. **Escalate before guessing** — wrong actions at 2 AM cost more than a call\n\n## MOP Structure\n\n| Section | Content |\n|---|---|\n| Objective | What the change accomplishes |\n| Prerequisites | Equipment state, access, backups |\n| Procedure | Numbered steps + expected outputs |\n| Verification | How to confirm success |\n| Rollback | How to undo on failure |\n\n## NOC Escalation Tiers\n\n- **Tier 1**: Alarm monitoring, initial triage (<15 min)\n- **Tier 2**: Optical engineer, CLI access (15–60 min)\n- **Tier 3**: Architect + vendor TAC (>60 min or SLA at risk)"},
        {"title": "Ch 2 — Performance Monitoring Mastery", "slug": "ch2-pm-mastery", "content_md": "# Ch 2 — Performance Monitoring Mastery\n\n| Layer | Parameters | Intervals |\n|---|---|---|\n| Optical | OPR, OPT, OSNR, ORL | 15-min, 24-hr |\n| OTN | BER pre/post-FEC, ES, SES | 15-min, 24-hr |\n| Client | Frame errors, CRC | 15-min |\n\n## TCA Design\n\n- Pre-FEC BER TCA: 1E-5 (7 dB margin from FEC threshold)\n- OSNR TCA: required + 2 dB\n- OPR TCA: ±2 dB from commissioned baseline\n\n## Case Study: Failing EDFA Pump\n\nSymptom: All-channel Q-factor down 0.5 dB over 72 hr. Root cause: 980 nm pump diode aging. Action: Increase pump drive current (temporary) → schedule EDFA replacement."},
        {"title": "Ch 3 — Alarm Management & Triage", "slug": "ch3-alarm-triage", "content_md": "# Ch 3 — Alarm Management & Triage\n\n## Classification Rules\n\n1. **Root cause**: Alarm on the failing object itself\n2. **Symptom**: Downstream alarms caused by root cause\n3. **Transient**: Clears within 30 sec — log, don't escalate\n\n## Correlation\n\n```\nMultiple alarms?\n→ Common upstream point → site power/fiber entry\n→ All on one fiber → physical fiber event\n→ Scattered → check amplifier chain\n→ Single card → replace hardware\n```\n\n## 96-Channel Fiber Cut\n\n1 LOS generates up to 290 alarms (96 OTU-LOF + 96 ODU-AIS + upstream). Triage: acknowledge secondary alarms, focus on root LOS, dispatch fiber crew."},
        {"title": "Ch 4 — Multi-Vendor CLI Reference", "slug": "ch4-multi-vendor-cli", "content_md": "# Ch 4 — Multi-Vendor CLI Reference\n\n## Ciena WaveLogic / SAOS\n\n```bash\n# Show all channels\nwaveserver channels show\n# PM data\nwaveserver pm get channel-index 1 interval current\n# Amplifier status\nwaveserver amplifiers show\n# Set target power\nwaveserver channels set channel-index 1 target-power -3.0\n```\n\n## Nokia 1830 PSS\n\n```bash\nshow port-type optical\nshow amplifier gain-tilt\nshow pm optical port-id 1-1-1 interval 15-min count 4\nclear pm tca port-id 1-1-1\n```\n\n## Cisco NCS — IOS-XR\n\n```\nshow controllers optics 0/0/0/0\nshow controllers optics 0/0/0/0 pm current 15-min optics 1\nshow controllers coherentDSP 0/0/0/0\n```\n\n## Infinera TNMS\n\n```bash\nshow optical line-system\nshow optical channel power all\nshow fec pm current\n```"},
        {"title": "Ch 5 — Commissioning Procedures", "slug": "ch5-commissioning", "content_md": "# Ch 5 — Commissioning Procedures\n\n## Pre-Commissioning Checklist\n\n- [ ] OTDR traces reviewed, no events > 0.5 dB\n- [ ] Connectors inspected and clean\n- [ ] Site power: -48V ±5%\n- [ ] NMS connectivity confirmed to all NEs\n- [ ] Channels in NMS match physical fiber map\n\n## Sequence\n\n1. Launch OSC — verify span connectivity\n2. Launch probe at -10 dBm — measure loss\n3. Set EDFA gain, enable channels\n4. Verify per-channel OSNR\n5. Log baseline: OPR, OPT, OSNR, Q per channel\n\n## Acceptance Criteria\n\n| Test | Pass |\n|---|---|\n| Channel OSNR | ≥ required + 3 dB |\n| Pre-FEC BER | < 1E-4 |\n| Protection switch | < 50 ms |"},
        {"title": "Ch 6 — Systematic Troubleshooting", "slug": "ch6-troubleshooting", "content_md": "# Ch 6 — Systematic Troubleshooting\n\n## 6-Phase Framework\n\n1. **OBSERVE** — Collect alarms, PM snapshot, recent changes\n2. **LOCALIZE** — Identify failing span or element\n3. **ISOLATE** — Narrow to root cause using half-split\n4. **DIAGNOSE** — Confirm with measurements\n5. **REMEDIATE** — Fix per approved MOP\n6. **DOCUMENT** — Close ticket, file RCA, update baseline\n\n## Half-Split Technique\n\nFor 6-span degraded OSNR: check span 3 output first → binary search finds fault in ≤ log₂(N) steps.\n\n## 5-Why RCA\n\nSymptom: Rx power -4 dBm low → EDFA output low → gain target changed at maintenance → engineer changed without updating as-built → MOP missing baseline verification step → Fix: update MOP template."},
        {"title": "Ch 7 — Fault Case Studies", "slug": "ch7-fault-case-studies", "content_md": "# Ch 7 — Fault Case Studies\n\n## Case 1: EDFA Pump Degradation\n\nSymptom: All channels Q-factor trending down 0.1 dB/week. Root cause: 980 nm pump diode aging. Resolution: EDFA replacement in maintenance window.\n\n## Case 2: Thermal BER Spikes\n\nSymptom: BER spikes daily 14:00–17:00 PDT. Root cause: Relay shelter ambient >45°C causes DSP clock jitter. Resolution: HVAC repair + temperature TCA.\n\n## Case 3: PMD-Induced Lock Loss\n\nSymptom: One channel drops every 72–96 hours. Root cause: G.653 segment PMD = 42 ps (limit: 30 ps). Resolution: Replace 15 km with G.654E.\n\n## Case 4: NMS-NE Sync Failure\n\nSymptom: NMS shows channels 'unknown', traffic passing. Root cause: NE management CPU 95% — PM storm. Resolution: Reduce polling interval 30s → 5 min."},
        {"title": "Lab 3.8 — War Room: ApexCarrier Emergency", "slug": "lab-3-8-war-room", "content_md": "# Lab 3.8 — War Room: ApexCarrier Emergency\n\n**Time**: 02:17 AM. 112 alarms. Three simultaneous events.\n\n## Event 1: SEA–PDX Fiber Cut\n\nLOS on ROADM-SEA port 5 Rx at 01:58. 48 OTU-LOF, 48 ODU-AIS. SLA: 99.999%.\n\n## Event 2: DEN EDFA Oscillation\n\nEDFA-DEN-3 gain tilt ±4 dB (normal ±0.5). 6 channels Q < 10 dB.\n\n## Event 3: ORD NMS Connectivity Loss\n\nNETCONF lost to all 6 Chicago NEs. Traffic status unknown. Last change: software upgrade 3 hours ago.\n\n## Your Deliverables\n\n1. Triage sequence with justification\n2. Exact CLI for Event 1 on Ciena ROADM-SEA\n3. 60-second customer bridge script\n4. MOP deviation + approval chain documentation\n5. 5-section RCA outline"}
    ]
}

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--docs-dir', default=str(Path.home() / 'Downloads'))
    args = parser.parse_args()
    docs_dir = Path(args.docs_dir)

    catalog = {}
    docx_files = list(docs_dir.glob('DWDM Architect Program*.docx'))
    if not docx_files:
        docx_files = list(docs_dir.glob('DWDM*.docx'))

    if docx_files:
        print(f"Found {len(docx_files)} .docx file(s) in {docs_dir}")
        for f in sorted(docx_files):
            name = f.name
            print(f"  Processing: {name}")
            level = None
            name_lower = name.lower()
            if 'level 1' in name_lower or 'foundations' in name_lower:
                level = 1
            elif 'level 2' in name_lower or 'architecture' in name_lower:
                level = 2
            elif 'level 3' in name_lower or 'production' in name_lower or 'troubleshooting' in name_lower:
                level = 3
            if level is None:
                print(f"    Skipping — cannot determine level")
                continue
            try:
                secs = docx_to_sections(f)
                if level not in catalog:
                    catalog[level] = []
                seen = {s['slug'] for s in catalog[level]}
                added = 0
                for sec in secs:
                    if sec['slug'] not in seen:
                        catalog[level].append(sec)
                        seen.add(sec['slug'])
                        added += 1
                print(f"    → Level {level}: {added} sections extracted")
            except Exception as e:
                print(f"    ✗ Error: {e}")
    else:
        print(f"No .docx files found in {docs_dir}")

    for lv in [1, 2, 3]:
        if lv not in catalog or not catalog[lv]:
            print(f"  Using embedded fallback for Level {lv}")
            catalog[lv] = FALLBACK[lv]

    out_json = CONTENT_DIR / 'catalog.json'
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(catalog, f, indent=2, ensure_ascii=False)

    out_js = CONTENT_DIR / 'catalog.js'
    with open(out_js, 'w', encoding='utf-8') as f:
        f.write('// Auto-generated by scripts/extract_catalog.py — do not edit.\n\n')
        f.write('const CATALOG = ')
        json.dump(catalog, f, indent=2, ensure_ascii=False)
        f.write(';\n\nexport default CATALOG;\n')

    total = sum(len(v) for v in catalog.values())
    print(f"\n✅ {total} sections written across {len(catalog)} levels")
    print(f"   {out_json} ({out_json.stat().st_size // 1024} KB)")
    print(f"   {out_js}  ({out_js.stat().st_size // 1024} KB)")
    print('\nNext: npm run db:seed')

if __name__ == '__main__':
    main()
